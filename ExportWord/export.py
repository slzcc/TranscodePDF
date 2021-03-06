#!/usr/bin/env python
# -*- coding:utf-8 -*-
from django.http import JsonResponse, HttpResponse
from django import conf
import datetime, re, os, time, requests, json, uuid
import docx
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def WriteFile(ResumeCandidateInfo, filePath):

	fileName = str(uuid.uuid1()) + ".docx"

	newWordFilePath = os.path.join(conf.settings.BASE_DIR, "static/storage", fileName)

	if not os.path.isfile(newWordFilePath):
		doc = Document(filePath)
		doc.styles['Normal'].font.name = '微软雅黑'
		doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

		paragraph = doc.add_paragraph()
		paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		run = paragraph.add_run('个人简历')
		run.bold = True
		run.font.name=u'微软雅黑'
		run.font.size=Pt(11)

		# 换行
		doc.add_paragraph()

		BaseInfo = {}
		OtherData = {}

		# 初始化数据
		for k, v in ResumeCandidateInfo.items():
			if k == "BaseInfo":
				# 写入基础数据
				if BaseInfo:
					paragraph = doc.add_paragraph()
					run = paragraph.add_run(u"基础信息")
					run.font.name='微软雅黑'
					run.font.size=Pt(11)
					run.bold = True
				for vs in sorted(v.items()):
					paragraph = doc.add_paragraph()
					base_key, base_value = vs[1].split("|")[0], vs[1].split("|")[1]
					run = paragraph.add_run(base_key + ": " + str(base_value) if base_value else "")
					run.font.name='微软雅黑'
					run.font.size=Pt(10.5)
			else:
				OtherData[k] = v

		# 换行
		doc.add_paragraph()

		# 写入其他数据
		for k, v in OtherData.items():

			paragraph = doc.add_paragraph()
			run = paragraph.add_run(k)
			run.bold = True
			run.font.name='微软雅黑'
			run.font.size=Pt(11)

			paragraph = doc.add_paragraph()
			_Data = v.replace("\r", '').replace("\n\n", '\n')
			run = paragraph.add_run(_Data)
			run.font.name='微软雅黑'
			run.font.size=Pt(10.5)

			doc.add_paragraph()

		# 保存文件
		doc.save(newWordFilePath)

	return newWordFilePath

def GenerateWordDocument(request):
	ret = {"status": "seccuss", "status_code": "200"}

	return JsonResponse(ret)